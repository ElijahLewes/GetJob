from typing import Dict
from docx import Document


def write_documents(
    job: Dict, resume_text: str, cover_letter_text: str, out_dir: str
) -> None:
    """Save customized resume and cover letter to Word documents."""
    doc = Document()
    doc.add_paragraph(resume_text)
    doc.save(f"{out_dir}/customized_resume.docx")

    cl = Document()
    cl.add_paragraph(cover_letter_text)
    cl.save(f"{out_dir}/cover_letter.docx")
